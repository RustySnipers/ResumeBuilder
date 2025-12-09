"""
DOCX Generator - Phase 5

Generate professional DOCX resumes compatible with Microsoft Word.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List, Optional
from io import BytesIO
import logging

logger = logging.getLogger(__name__)


class DOCXGenerator:
    """
    Generate professional ATS-friendly DOCX resumes.

    Compatible with Microsoft Word and other office suites.
    """

    def __init__(self):
        """Initialize DOCX generator."""
        self.doc = None

    def generate(
        self,
        resume_data: Dict[str, Any],
        template: str = "professional",
        output_path: Optional[str] = None,
    ) -> bytes:
        """
        Generate DOCX resume from resume data.

        Args:
            resume_data: Resume content as dictionary
            template: Template style to use
            output_path: Optional file path to save DOCX

        Returns:
            DOCX content as bytes
        """
        # Create new document
        self.doc = Document()

        # Set document properties
        core_props = self.doc.core_properties
        core_props.title = f"Resume - {resume_data.get('name', 'Unknown')}"
        core_props.author = resume_data.get('name', 'Unknown')
        core_props.subject = "Professional Resume"

        # Set up styles
        self._configure_styles(template)

        # Build content
        self._build_header(resume_data)
        self._add_spacing(0.1)

        if "summary" in resume_data and resume_data["summary"]:
            self._build_summary(resume_data["summary"])
            self._add_spacing(0.15)

        if "experience" in resume_data and resume_data["experience"]:
            self._build_experience(resume_data["experience"])
            self._add_spacing(0.15)

        if "education" in resume_data and resume_data["education"]:
            self._build_education(resume_data["education"])
            self._add_spacing(0.15)

        if "skills" in resume_data and resume_data["skills"]:
            self._build_skills(resume_data["skills"])
            self._add_spacing(0.15)

        if "certifications" in resume_data and resume_data["certifications"]:
            self._build_certifications(resume_data["certifications"])
            self._add_spacing(0.15)

        if "projects" in resume_data and resume_data["projects"]:
            self._build_projects(resume_data["projects"])

        # Save to buffer
        buffer = BytesIO()
        self.doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()

        # Save to file if path provided
        if output_path:
            self.doc.save(output_path)

        logger.info(f"Generated DOCX resume ({len(docx_bytes)} bytes)")

        return docx_bytes

    def _configure_styles(self, template: str = "standard_ats"):
        """Configure document styles based on template."""
        # Set narrow margins (0.5 inch all around)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Template Configurations
        if template == "modern_ats":
            self.header_color = RGBColor(44, 90, 160)  # Dark Blue
            self.text_color = RGBColor(33, 33, 33)     # Dark Grey
            self.sub_text_color = RGBColor(80, 80, 80) # Muted
            self.font_name = 'Calibri'
            self.title_size = Pt(24)
            self.h2_size = Pt(14)
            self.h2_caps = False
        elif template == "strict_ats":
            # strict_ats (iCIMS/Taleo optimized)
            self.header_color = RGBColor(0, 0, 0)      # Black
            self.text_color = RGBColor(0, 0, 0)        # Black
            self.sub_text_color = RGBColor(0, 0, 0)    # Black
            self.font_name = 'Arial'
            self.title_size = Pt(20)                   # Slightly smaller title
            self.h2_size = Pt(11)                      # Standard header size
            self.h2_caps = True                        # UPPERCASE HEADERS
        else:
            # standard_ats (Default)
            self.header_color = RGBColor(0, 0, 0)      # Black
            self.text_color = RGBColor(0, 0, 0)        # Black
            self.sub_text_color = RGBColor(0, 0, 0)    # Black
            self.font_name = 'Arial'
            self.title_size = Pt(22)
            self.h2_size = Pt(12)
            self.h2_caps = True

        # Modify existing styles
        styles = self.doc.styles
        
        # Normal style
        style_normal = styles['Normal']
        font_normal = style_normal.font
        font_normal.name = self.font_name
        font_normal.size = Pt(10.5)
        font_normal.color.rgb = self.text_color

        # Heading 1 (Name)
        style_h1 = styles['Heading 1']
        font_h1 = style_h1.font
        font_h1.name = self.font_name
        font_h1.size = self.title_size
        font_h1.bold = True
        font_h1.color.rgb = RGBColor(0, 0, 0) # Name always black

        # Heading 2 (Section headers)
        style_h2 = styles['Heading 2']
        font_h2 = style_h2.font
        font_h2.name = self.font_name
        font_h2.size = self.h2_size
        font_h2.bold = True
        font_h2.color.rgb = self.header_color
        # Note: All Caps handled in build method as python-docx style caps is tricky

        # Heading 3 (Job titles)
        style_h3 = styles['Heading 3']
        font_h3 = style_h3.font
        font_h3.name = self.font_name
        font_h3.size = Pt(11)
        font_h3.bold = True
        font_h3.color.rgb = self.text_color

    def _build_header(self, resume_data: Dict[str, Any]):
        """Build resume header with name and contact info."""
        # Name
        if "name" in resume_data:
            name_text = resume_data["name"]
            name_para = self.doc.add_paragraph(name_text, style='Heading 1')
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact information
        contact_parts = []
        if "email" in resume_data:
            contact_parts.append(resume_data["email"])
        if "phone" in resume_data:
            contact_parts.append(resume_data["phone"])
        if "location" in resume_data:
            contact_parts.append(resume_data["location"])
        if "linkedin" in resume_data:
            contact_parts.append(f"LinkedIn: {resume_data['linkedin']}")
        if "website" in resume_data:
            contact_parts.append(resume_data["website"])

        if contact_parts:
            contact_text = " | ".join(contact_parts)
            contact_para = self.doc.add_paragraph(contact_text)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = contact_para.runs[0]
            run.font.size = Pt(10)
            run.font.color.rgb = self.sub_text_color

    def _build_section_header(self, title: str):
        """Build section header with optional all-caps."""
        text = title.upper() if getattr(self, 'h2_caps', False) else title
        self.doc.add_paragraph(text, style='Heading 2')

    def _build_summary(self, summary: str):
        """Build professional summary section."""
        self._build_section_header("PROFESSIONAL SUMMARY")
        self.doc.add_paragraph(summary)

    def _build_experience(self, experience: List[Dict[str, Any]]):
        """Build work experience section."""
        self._build_section_header("WORK EXPERIENCE")

        for job in experience:
            # Job title
            if "title" in job:
                self.doc.add_paragraph(job["title"], style='Heading 3')

            # Company and location
            company_parts = []
            if "company" in job:
                company_parts.append(job["company"])
            if "location" in job:
                company_parts.append(job["location"])

            if company_parts:
                company_para = self.doc.add_paragraph(" - ".join(company_parts))
                run = company_para.runs[0]
                run.font.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = self.sub_text_color

            # Dates
            date_parts = []
            if "start_date" in job:
                date_parts.append(job["start_date"])
            if "end_date" in job:
                date_parts.append(job["end_date"])
            elif "current" in job and job["current"]:
                date_parts.append("Present")

            if date_parts:
                date_para = self.doc.add_paragraph(" - ".join(date_parts))
                run = date_para.runs[0]
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = self.sub_text_color

            # Description
            if "description" in job and job["description"]:
                self.doc.add_paragraph(job["description"])

            # Achievements/bullets
            if "achievements" in job and job["achievements"]:
                for achievement in job["achievements"]:
                    self.doc.add_paragraph(achievement, style='List Bullet')

            # Add spacing between jobs
            self._add_spacing(0.1)

    def _build_education(self, education: List[Dict[str, Any]]):
        """Build education section."""
        self._build_section_header("EDUCATION")

        for edu in education:
            # Degree
            if "degree" in edu:
                self.doc.add_paragraph(edu["degree"], style='Heading 3')

            # Institution
            if "institution" in edu:
                inst_para = self.doc.add_paragraph(edu["institution"])
                run = inst_para.runs[0]
                run.font.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = self.sub_text_color

            # Dates
            date_parts = []
            if "start_date" in edu:
                date_parts.append(edu["start_date"])
            if "end_date" in edu:
                date_parts.append(edu["end_date"])

            if date_parts:
                date_para = self.doc.add_paragraph(" - ".join(date_parts))
                run = date_para.runs[0]
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = self.sub_text_color

            # GPA
            if "gpa" in edu:
                self.doc.add_paragraph(f"GPA: {edu['gpa']}")

            # Honors/Awards
            if "honors" in edu and edu["honors"]:
                for honor in edu["honors"]:
                    self.doc.add_paragraph(honor, style='List Bullet')

            # Add spacing between education entries
            self._add_spacing(0.1)

    def _build_skills(self, skills: List[str]):
        """Build skills section."""
        self._build_section_header("SKILLS")
        skills_text = ", ".join(skills)
        self.doc.add_paragraph(skills_text)

    def _build_certifications(self, certifications: List[Dict[str, Any]]):
        """Build certifications section."""
        self._build_section_header("CERTIFICATIONS")

        for cert in certifications:
            # Certificate name
            if "name" in cert:
                self.doc.add_paragraph(cert["name"], style='Heading 3')

            # Issuer
            if "issuer" in cert:
                issuer_para = self.doc.add_paragraph(cert["issuer"])
                run = issuer_para.runs[0]
                run.font.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = self.sub_text_color

            # Date
            if "date" in cert:
                date_para = self.doc.add_paragraph(cert["date"])
                run = date_para.runs[0]
                run.font.size = Pt(10)
                run.font.italic = True

            # Credential ID
            if "credential_id" in cert:
                self.doc.add_paragraph(f"Credential ID: {cert['credential_id']}")

            # Add spacing between certifications
            self._add_spacing(0.1)

    def _build_projects(self, projects: List[Dict[str, Any]]):
        """Build projects section."""
        self._build_section_header("PROJECTS")

        for project in projects:
            # Project name
            if "name" in project:
                self.doc.add_paragraph(project["name"], style='Heading 3')

            # Description
            if "description" in project:
                self.doc.add_paragraph(project["description"])

            # Technologies
            if "technologies" in project:
                tech_text = f"Technologies: {', '.join(project['technologies'])}"
                tech_para = self.doc.add_paragraph(tech_text)
                run = tech_para.runs[0]
                run.font.italic = True
                run.font.color.rgb = self.sub_text_color

            # URL
            if "url" in project:
                self.doc.add_paragraph(f"URL: {project['url']}")

            # Add spacing between projects
            self._add_spacing(0.1)

    def _add_spacing(self, inches: float):
        """Add vertical spacing."""
        self.doc.add_paragraph()
        last_para = self.doc.paragraphs[-1]
        last_para.paragraph_format.space_after = Pt(inches * 72)

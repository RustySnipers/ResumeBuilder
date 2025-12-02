"""
Unit Tests for Phase 5 Export System

Tests for PDF generator, DOCX generator, and template engine.
"""

import pytest
from io import BytesIO
from backend.export.pdf_generator import PDFGenerator
from backend.export.docx_generator import DOCXGenerator
from backend.export.template_engine import TemplateEngine
from PyPDF2 import PdfReader
from docx import Document


# ============================================================================
# Test Data
# ============================================================================

SAMPLE_RESUME_DATA = {
    "name": "John Doe",
    "email": "john.doe@example.com",
    "phone": "+1 (555) 123-4567",
    "location": "San Francisco, CA",
    "linkedin": "linkedin.com/in/johndoe",
    "website": "johndoe.com",
    "summary": "Experienced software engineer with 5+ years of experience in full-stack development.",
    "experience": [
        {
            "title": "Senior Software Engineer",
            "company": "Tech Corp",
            "location": "San Francisco, CA",
            "start_date": "Jan 2020",
            "end_date": "Present",
            "current": True,
            "description": "Led development of microservices architecture",
            "achievements": [
                "Increased system performance by 50%",
                "Mentored 5 junior engineers"
            ]
        },
        {
            "title": "Software Engineer",
            "company": "StartUp Inc",
            "location": "San Jose, CA",
            "start_date": "Jun 2018",
            "end_date": "Dec 2019",
            "current": False,
            "description": "Developed web applications using React and Node.js",
            "achievements": [
                "Built 3 customer-facing products",
                "Reduced load time by 40%"
            ]
        }
    ],
    "education": [
        {
            "degree": "Bachelor of Science in Computer Science",
            "institution": "University of California, Berkeley",
            "start_date": "2014",
            "end_date": "2018",
            "gpa": "3.8",
            "honors": ["Dean's List", "Summa Cum Laude"]
        }
    ],
    "skills": ["Python", "JavaScript", "React", "Node.js", "AWS", "Docker", "Kubernetes"],
    "certifications": [
        {
            "name": "AWS Certified Solutions Architect",
            "issuer": "Amazon Web Services",
            "date": "2023",
            "credential_id": "ABC123XYZ"
        }
    ],
    "projects": [
        {
            "name": "Open Source Project",
            "description": "Built a tool for automated testing",
            "technologies": ["Python", "Docker", "CI/CD"],
            "url": "github.com/johndoe/project"
        }
    ]
}

MINIMAL_RESUME_DATA = {
    "name": "Jane Smith",
    "email": "jane@example.com"
}


# ============================================================================
# PDF Generator Tests
# ============================================================================

class TestPDFGenerator:
    """Test suite for PDF generation."""

    def test_pdf_generator_initialization(self):
        """Test PDF generator can be initialized."""
        generator = PDFGenerator()
        assert generator is not None
        assert generator.font_name == "Helvetica"

    def test_generate_pdf_with_full_data(self):
        """Test PDF generation with complete resume data."""
        generator = PDFGenerator()
        pdf_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        assert pdf_bytes is not None
        assert len(pdf_bytes) > 0
        assert pdf_bytes[:4] == b'%PDF'  # PDF magic number

    def test_generate_pdf_with_minimal_data(self):
        """Test PDF generation with minimal resume data."""
        generator = PDFGenerator()
        pdf_bytes = generator.generate(MINIMAL_RESUME_DATA, template="professional")

        assert pdf_bytes is not None
        assert len(pdf_bytes) > 0
        assert pdf_bytes[:4] == b'%PDF'

    def test_pdf_contains_user_name(self):
        """Test that PDF contains the user's name."""
        generator = PDFGenerator()
        pdf_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        # Read PDF and check for name
        pdf_reader = PdfReader(BytesIO(pdf_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()

        assert "John Doe" in text

    def test_pdf_contains_experience(self):
        """Test that PDF contains experience section."""
        generator = PDFGenerator()
        pdf_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        pdf_reader = PdfReader(BytesIO(pdf_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()

        assert "Senior Software Engineer" in text
        assert "Tech Corp" in text

    def test_pdf_contains_education(self):
        """Test that PDF contains education section."""
        generator = PDFGenerator()
        pdf_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        pdf_reader = PdfReader(BytesIO(pdf_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()

        assert "Computer Science" in text
        assert "University of California" in text

    def test_pdf_contains_skills(self):
        """Test that PDF contains skills section."""
        generator = PDFGenerator()
        pdf_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        pdf_reader = PdfReader(BytesIO(pdf_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()

        assert "Python" in text or "JavaScript" in text

    def test_pdf_metadata(self):
        """Test that PDF has proper metadata."""
        generator = PDFGenerator()
        pdf_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        pdf_reader = PdfReader(BytesIO(pdf_bytes))
        metadata = pdf_reader.metadata

        assert metadata is not None
        assert "/Title" in metadata or "title" in str(metadata).lower()

    def test_pdf_generation_performance(self):
        """Test that PDF generation completes in reasonable time."""
        import time

        generator = PDFGenerator()
        start_time = time.time()
        pdf_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")
        end_time = time.time()

        generation_time = end_time - start_time
        assert generation_time < 5.0  # Should complete in under 5 seconds

    def test_pdf_with_empty_sections(self):
        """Test PDF generation with empty sections."""
        data = {
            "name": "Test User",
            "email": "test@example.com",
            "experience": [],
            "education": [],
            "skills": []
        }

        generator = PDFGenerator()
        pdf_bytes = generator.generate(data, template="professional")

        assert pdf_bytes is not None
        assert len(pdf_bytes) > 0


# ============================================================================
# DOCX Generator Tests
# ============================================================================

class TestDOCXGenerator:
    """Test suite for DOCX generation."""

    def test_docx_generator_initialization(self):
        """Test DOCX generator can be initialized."""
        generator = DOCXGenerator()
        assert generator is not None

    def test_generate_docx_with_full_data(self):
        """Test DOCX generation with complete resume data."""
        generator = DOCXGenerator()
        docx_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        assert docx_bytes is not None
        assert len(docx_bytes) > 0
        # DOCX files start with PK (ZIP format)
        assert docx_bytes[:2] == b'PK'

    def test_generate_docx_with_minimal_data(self):
        """Test DOCX generation with minimal resume data."""
        generator = DOCXGenerator()
        docx_bytes = generator.generate(MINIMAL_RESUME_DATA, template="professional")

        assert docx_bytes is not None
        assert len(docx_bytes) > 0

    def test_docx_contains_user_name(self):
        """Test that DOCX contains the user's name."""
        generator = DOCXGenerator()
        docx_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        # Read DOCX
        doc = Document(BytesIO(docx_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])

        assert "John Doe" in text

    def test_docx_contains_experience(self):
        """Test that DOCX contains experience section."""
        generator = DOCXGenerator()
        docx_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        doc = Document(BytesIO(docx_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])

        assert "Senior Software Engineer" in text
        assert "Tech Corp" in text

    def test_docx_contains_education(self):
        """Test that DOCX contains education section."""
        generator = DOCXGenerator()
        docx_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        doc = Document(BytesIO(docx_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])

        assert "Computer Science" in text

    def test_docx_contains_skills(self):
        """Test that DOCX contains skills section."""
        generator = DOCXGenerator()
        docx_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        doc = Document(BytesIO(docx_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])

        assert "Python" in text or "JavaScript" in text

    def test_docx_metadata(self):
        """Test that DOCX has proper metadata."""
        generator = DOCXGenerator()
        docx_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")

        doc = Document(BytesIO(docx_bytes))
        core_props = doc.core_properties

        assert core_props.title is not None
        assert "John Doe" in core_props.title or "Resume" in core_props.title

    def test_docx_generation_performance(self):
        """Test that DOCX generation completes in reasonable time."""
        import time

        generator = DOCXGenerator()
        start_time = time.time()
        docx_bytes = generator.generate(SAMPLE_RESUME_DATA, template="professional")
        end_time = time.time()

        generation_time = end_time - start_time
        assert generation_time < 3.0  # Should complete in under 3 seconds

    def test_docx_with_empty_sections(self):
        """Test DOCX generation with empty sections."""
        data = {
            "name": "Test User",
            "email": "test@example.com",
            "experience": [],
            "education": [],
            "skills": []
        }

        generator = DOCXGenerator()
        docx_bytes = generator.generate(data, template="professional")

        assert docx_bytes is not None
        assert len(docx_bytes) > 0


# ============================================================================
# Template Engine Tests
# ============================================================================

class TestTemplateEngine:
    """Test suite for template engine."""

    def test_template_engine_initialization(self):
        """Test template engine can be initialized."""
        engine = TemplateEngine()
        assert engine is not None
        assert engine.templates_dir is not None

    def test_list_templates(self):
        """Test listing available templates."""
        engine = TemplateEngine()
        templates = engine.list_templates()

        assert templates is not None
        assert len(templates) > 0
        assert isinstance(templates, list)

        # Check template structure
        for template in templates:
            assert "id" in template
            assert "name" in template
            assert "description" in template
            assert "category" in template

    def test_get_template_info_professional(self):
        """Test getting professional template info."""
        engine = TemplateEngine()
        template_info = engine.get_template_info("professional")

        assert template_info is not None
        assert template_info["id"] == "professional"
        assert template_info["name"] == "Professional"

    def test_get_template_info_modern(self):
        """Test getting modern template info."""
        engine = TemplateEngine()
        template_info = engine.get_template_info("modern")

        assert template_info is not None
        assert template_info["id"] == "modern"
        assert template_info["name"] == "Modern Professional"

    def test_get_template_info_invalid(self):
        """Test getting info for non-existent template."""
        engine = TemplateEngine()

        with pytest.raises(ValueError):
            engine.get_template_info("nonexistent")

    def test_render_template_with_data(self):
        """Test rendering template with resume data."""
        engine = TemplateEngine()
        html = engine.render("professional", SAMPLE_RESUME_DATA)

        assert html is not None
        assert len(html) > 0
        assert "John Doe" in html
        assert "john.doe@example.com" in html

    def test_render_template_contains_structure(self):
        """Test that rendered HTML has proper structure."""
        engine = TemplateEngine()
        html = engine.render("professional", SAMPLE_RESUME_DATA)

        assert "<!DOCTYPE html>" in html
        assert "<html" in html
        assert "</html>" in html
        assert "<head>" in html
        assert "<body>" in html

    def test_format_phone_filter(self):
        """Test phone number formatting filter."""
        engine = TemplateEngine()

        # Test US phone number formatting
        formatted = engine._format_phone("5551234567")
        assert formatted == "(555) 123-4567"

        # Test with country code
        formatted = engine._format_phone("15551234567")
        assert "+1" in formatted

    def test_template_render_performance(self):
        """Test that template rendering completes quickly."""
        import time

        engine = TemplateEngine()
        start_time = time.time()
        html = engine.render("professional", SAMPLE_RESUME_DATA)
        end_time = time.time()

        render_time = end_time - start_time
        assert render_time < 1.0  # Should complete in under 1 second


# ============================================================================
# Integration Tests (Generators + Data)
# ============================================================================

class TestExportIntegration:
    """Integration tests for export system."""

    def test_pdf_and_docx_consistency(self):
        """Test that PDF and DOCX contain same information."""
        pdf_gen = PDFGenerator()
        docx_gen = DOCXGenerator()

        pdf_bytes = pdf_gen.generate(SAMPLE_RESUME_DATA, template="professional")
        docx_bytes = docx_gen.generate(SAMPLE_RESUME_DATA, template="professional")

        # Both should be generated
        assert pdf_bytes is not None and len(pdf_bytes) > 0
        assert docx_bytes is not None and len(docx_bytes) > 0

        # Extract text from both
        pdf_reader = PdfReader(BytesIO(pdf_bytes))
        pdf_text = ""
        for page in pdf_reader.pages:
            pdf_text += page.extract_text()

        doc = Document(BytesIO(docx_bytes))
        docx_text = "\n".join([para.text for para in doc.paragraphs])

        # Check both contain key information
        assert "John Doe" in pdf_text and "John Doe" in docx_text
        assert "Tech Corp" in pdf_text and "Tech Corp" in docx_text

    def test_special_characters_handling(self):
        """Test handling of special characters in resume data."""
        data = {
            "name": "José García-Müller",
            "email": "josé@example.com",
            "summary": "Expert in C++, C# & Python. Has 10+ years' experience.",
            "skills": ["C++", "C#", "Python", "R&D"]
        }

        pdf_gen = PDFGenerator()
        docx_gen = DOCXGenerator()

        # Should not raise exceptions
        pdf_bytes = pdf_gen.generate(data, template="professional")
        docx_bytes = docx_gen.generate(data, template="professional")

        assert pdf_bytes is not None
        assert docx_bytes is not None

    def test_long_resume_handling(self):
        """Test handling of long resumes (multiple pages)."""
        # Create resume with lots of experience
        long_data = SAMPLE_RESUME_DATA.copy()
        long_data["experience"] = [
            {
                "title": f"Position {i}",
                "company": f"Company {i}",
                "start_date": "2020",
                "end_date": "2021",
                "description": "Long description " * 20,
                "achievements": [f"Achievement {j}" for j in range(5)]
            }
            for i in range(10)
        ]

        pdf_gen = PDFGenerator()
        docx_gen = DOCXGenerator()

        # Should handle multi-page resumes
        pdf_bytes = pdf_gen.generate(long_data, template="professional")
        docx_bytes = docx_gen.generate(long_data, template="professional")

        assert pdf_bytes is not None
        assert docx_bytes is not None


if __name__ == "__main__":
    pytest.main([__file__, "-v"])

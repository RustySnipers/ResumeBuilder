import requests
import time
import os
import sys
from docx import Document

BASE_URL = "http://127.0.0.1:8000"

def print_pass(msg):
    print(f"[PASS] {msg}")

def print_fail(msg):
    print(f"[FAIL] {msg}")
    sys.exit(1)

def create_test_docx(filename="test_resume.docx"):
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('john.doe@example.com | 123-456-7890')
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Experienced Python Developer with expertise in FastAPI and AI.')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Senior Developer at Tech Co.')
    doc.save(filename)
    return filename

def run_tests():
    # 1. Health Check
    try:
        r = requests.get(f"{BASE_URL}/health/live")
        if r.status_code == 200:
            print_pass("Health Check (Live)")
        else:
            print_fail(f"Health Check failed: {r.status_code}")
    except Exception as e:
        print_fail(f"Could not connect to server: {e}")

    # 2. Detailed Health
    r = requests.get(f"{BASE_URL}/health")
    data = r.json()
    print(f"    Services Status: {data.get('services')}")
    print_pass("Detailed Health Check")

    # 3. Auth (Register & Login)
    token = None
    email = f"test_{int(time.time())}@example.com"
    password = "Password123!"
    
    try:
        # Register
        reg_payload = {"email": email, "password": password, "full_name": "Test User"}
        r = requests.post(f"{BASE_URL}/auth/register", json=reg_payload)
        if r.status_code == 201:
            print_pass("User Registration")
        else:
            print_fail(f"Registration failed: {r.text}")

        # Login
        login_data = {"username": email, "password": password}
        r = requests.post(f"{BASE_URL}/auth/login", data=login_data)
        if r.status_code == 200:
            token = r.json().get("access_token")
            print_pass("User Login")
        else:
            print_fail(f"Login failed: {r.text}")
            
    except Exception as e:
        print_fail(f"Auth Exception: {e}")

    headers = {"Authorization": f"Bearer {token}"}

    # 4. Create & Upload Resume
    filename = create_test_docx()
    try:
        with open(filename, 'rb') as f:
            files = {'file': (filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            r = requests.post(
                f"{BASE_URL}/api/v1/resumes/upload", 
                files=files, 
                data={"title": "Test Resume"},
                headers=headers
            )
            
        if r.status_code == 201:
            print_pass("Resume Upload (DOCX)")
            resume_data = r.json()
            resume_id = resume_data['id']
            print(f"    Resume ID: {resume_id}")
            print(f"    Parsed Length: {len(resume_data.get('raw_text', ''))} chars")
        else:
            print_fail(f"Resume Upload failed: {r.status_code} - {r.text}")
    except Exception as e:
        print_fail(f"Resume Upload Exception: {e}")

    # 5. Analyze Job Fit
    job_desc = "Looking for a Senior Python Developer with FastAPI and AI experience."
    try:
        payload = {
            "resume_id": resume_id,
            "job_description": job_desc
        }
        print("    Requesting Analysis (may take time for model download)...")
        r = requests.post(
            f"{BASE_URL}/api/v1/jobs/analyze", 
            json=payload, 
            headers=headers,
            timeout=300 # Allow 5 minutes for model download/load
        )
        
        response = requests.post(f"{BASE_URL}/api/v1/export/docx", json=export_payload, headers=headers)
        if response.status_code == 200:
            print(f"    ✅ Export Successful. Size: {len(response.content)} bytes")
        else:
            print(f"    ❌ Export Failed: {response.text}")
            sys.exit(1)

        print("\n✅✅ FULL VERIFICATION SUCCESSFUL! ✅✅")

    except requests.exceptions.ConnectionError:
        print("\n❌ Could not connect to server. Is it running?")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ Verification Failed: {e}")
        sys.exit(1)
    finally:
        # Cleanup
        if os.path.exists(filename):
            os.remove(filename)

if __name__ == "__main__":
    print("Running Verification Tests...")
    # Wait a bit for server to fully settle if just launched
    time.sleep(2)
    run_tests()
